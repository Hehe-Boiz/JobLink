from __future__ import annotations

from abc import ABC, abstractmethod
from collections.abc import Iterator
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from urllib.parse import unquote, urlparse
from zipfile import BadZipFile, ZipFile
from .ocr import PDFOCRService

import pymupdf
import requests
from docx import Document
from docx.document import Document as DocxDocument
from docx.table import Table
from docx.text.paragraph import Paragraph

from apps.matching.constants import MAX_DOCUMENT_BYTES, MAX_DOCUMENT_PAGES, MAX_EXTRACTED_TEXT_CHARS

from .exceptions import DocumentDownloadError, EmptyDocumentError, InvalidDocumentError, UnsupportedDocumentError

PDF_EXTENSION = ".pdf"
DOCX_EXTENSION = ".docx"

PDF_CONTENT_TYPE = "application/pdf"

DOCX_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument."
    "wordprocessingml.document"
)

DOWNLOAD_CHUNK_SIZE = 64 * 1024

# DOCX là file ZIP nén.
# File upload 10 MB không nên giải nén thành hàng trăm MB.
MAX_DOCX_UNCOMPRESSED_BYTES = (
    MAX_DOCUMENT_BYTES * 20
)

@dataclass(frozen=True)
class ParsedDocument:
    content: bytes #-> Bytes gốc, dùng để tính cv_hash.
    text: str # lấy từ file bởi file text.py
    extension: str


class DocumentParser(ABC):
    @abstractmethod
    def parse(self, content: bytes) -> str:
        raise NotImplementedError

class PDFDocumentParser(DocumentParser):
    def __init__(self, ocr_service: PDFOCRService | None = None) -> None:
        self._ocr_service = (
            ocr_service or PDFOCRService()
        )

    def parse(self, content):
        if not content.startswith(b"%PDF-"):
            raise InvalidDocumentError("File không phải PDF hợp lệ.")

        document = None
        try:
            document = pymupdf.open(stream=content, filetype="pdf")
            if document.needs_pass:
                raise InvalidDocumentError("PDF được bảo vệ bằng mật khẩu.")

            if document.page_count > MAX_DOCUMENT_PAGES:
                raise InvalidDocumentError(f"PDF vượt quá giới hạn {MAX_DOCUMENT_PAGES} trang.")

            if document.page_count <= 0:
                raise EmptyDocumentError("PDF không có trang nào.")
            pages: list[str] = []
            total_text_chars = 0

            for page in document:
                page_text = page.get_text("text")
                total_text_chars += len(page_text)

                if (total_text_chars > MAX_EXTRACTED_TEXT_CHARS):
                    raise InvalidDocumentError("Nội dung CV vượt quá giới hạn xử lý.")

                pages.append(page_text)

        except (InvalidDocumentError, EmptyDocumentError,):
            raise



        except Exception as exc:
            raise InvalidDocumentError("File PDF không hợp lệ hoặc đã bị hỏng.") from exc


        finally:
            if document is not None:
                document.close()

        text = "\n".join(pages).strip()

        if text:
            return text

        return self._ocr_service.extract_text(content)

class DOCXDocumentParser(DocumentParser):
    def parse(self, content: bytes):
        self._validate_docx_archive(content)

        try:
            document = Document(BytesIO(content))

        except Exception as exc:
            raise InvalidDocumentError("File DOCX không hợp lệ hoặc đã bị hỏng.") from exc

        blocks = []
        total_text_chars = 0
        for block in iter_docx_blocks(document):

            if isinstance(block, Paragraph):

                text = block.text.strip()

                if not text:
                    continue

                blocks.append(text)

                total_text_chars += len(text)

            elif isinstance(block, Table):
                for row in block.rows:
                    cells = []
                    for cell in row.cells:
                        cell_text = (" ".join(cell.text.split()))

                        if cell_text:
                            cells.append(cell_text)

                    if not cells:
                        continue

                    # Giữ các cell trong cùng một row.
                    #
                    # Ví dụ:
                    #
                    # Python | Django | PostgreSQL
                    row_text = " | ".join(cells)

                    blocks.append(row_text)

                    total_text_chars += len(row_text)

            if (total_text_chars > MAX_EXTRACTED_TEXT_CHARS):
                raise InvalidDocumentError("Nội dung CV vượt quá giới hạn xử lý.")

        text = "\n".join(blocks).strip()

        if not text:
            raise EmptyDocumentError("Không đọc được nội dung từ file DOCX.")

        return text

    @staticmethod
    def _validate_docx_archive(content: bytes) -> None:
        """
        DOCX thực chất là ZIP.

        kiểm tra: -> Có phải ZIP không -> Có cấu trúc DOCX không -> Sau giải nén có quá lớn không
        """

        # ZIP thường bắt đầu bằng PK.
        if not content.startswith(b"PK"):
            raise InvalidDocumentError("File không phải DOCX hợp lệ.")

        try:
            # BytesIO(content) -> biến đống bytes đó thành một object giống file trong RAM, để ZipFile có thể đọc mà không cần lưu tạm xuống ổ cứng.
            with ZipFile(BytesIO(content)) as archive: # -> mở file DOCX như một ZIP archive.

                filenames = set(archive.namelist()) # -> rchive.namelist() trả về danh sách tên tất cả file bên trong ZIP.

                # Hai file quan trọng cho thấy đây
                # thực sự là Office Open XML document.
                required_files = {
                    "[Content_Types].xml",
                    "word/document.xml",
                }

                if not required_files.issubset(filenames): # kiểm tra xem có phải tập con không ?
                    raise InvalidDocumentError("File ZIP không phải DOCX hợp lệ.")

                # Tổng kích thước tất cả file
                # sau khi giải nén.
                total_uncompressed_bytes = 0

                file_infos = archive.infolist() # -> archive.infolist() trả về thông tin của từng file bên trong ZIP.

                for file_info in file_infos:
                    total_uncompressed_bytes = total_uncompressed_bytes + file_info.file_size

                if (total_uncompressed_bytes> MAX_DOCX_UNCOMPRESSED_BYTES):
                    raise InvalidDocumentError("DOCX sau khi giải nén vượt quá giới hạn an toàn.")

        except InvalidDocumentError:
            raise

        except BadZipFile as exc:
            raise InvalidDocumentError("File DOCX không hợp lệ hoặc đã bị hỏng.") from exc

def iter_docx_blocks(document: DocxDocument) -> Iterator[Paragraph | Table]:

    for child in (document.element.body.iterchildren()): # -> lấy phần <body> bên trong XML của Word

        if child.tag.endswith("}p"):
            yield Paragraph(child,document)

        elif child.tag.endswith("}tbl"):
            yield Table(child, document)


class CVDocumentLoader:
    def __init__(self) -> None:
        self._parsers: dict[str,DocumentParser] = {
            PDF_EXTENSION: (PDFDocumentParser()),
            DOCX_EXTENSION: (DOCXDocumentParser()),
        }

    def load(self, cv_field) -> ParsedDocument:
        if not cv_field:
            raise InvalidDocumentError("Application không có CV.")

        try:
            url = cv_field.url
        except Exception as exc:
            raise InvalidDocumentError("Không lấy được URL của CV.") from exc

        content, content_type = (self._download(url))

        extension = self._detect_extension(
            cv_field=cv_field,
            url=url,
            content_type=content_type, 
        )

        parser = self._parsers.get(extension)

        if parser is None:
            raise UnsupportedDocumentError(f"Chưa hỗ trợ định dạng CV: {extension or 'không xác định'}.")

        text = parser.parse(content)

        return ParsedDocument(
            content=content, # nếu có sẽ chứa toàn bộ nội dung file đã tải dưới dạng byte
            text=text,
            extension=extension,
        )

    def _detect_extension(self,cv_field, url: str, content_type: str) -> str:
        candidates = [str(cv_field), url]

        for candidate in candidates:
            path = unquote(urlparse(candidate).path)

            extension = (Path(path).suffix.lower())

            if extension in self._parsers:
                return extension

        content_type_map = {
            PDF_CONTENT_TYPE: PDF_EXTENSION,
            DOCX_CONTENT_TYPE: DOCX_EXTENSION,
        }

        extension = content_type_map.get(
            content_type
        )

        if extension:
            return extension

        raise UnsupportedDocumentError(
            "Không xác định được định dạng CV. "
            "Chỉ hỗ trợ PDF và DOCX."
        )

    def _download(self, url: str ) -> tuple[bytes, str]:
        try:
            with requests.get(url, stream=True, timeout=(5, 30)) as response: #-> stream để không tải hết 1 lúc 
                response.raise_for_status()
                content_type = (response.headers.get("Content-Type","").split(";", maxsplit=1)[0].strip().lower())

                content_length = (response.headers.get("Content-Length")) # lấy kích thước file 

                if content_length:
                    try:
                        declared_size = int(content_length)

                    except ValueError:
                        declared_size = None

                    if (declared_size is not None and declared_size > MAX_DOCUMENT_BYTES):
                        raise InvalidDocumentError("CV vượt quá giới hạn dung lượng cho phép.")

                downloaded = bytearray() #-> chỗ để chứa các chunk vừa tải

                for chunk in (response.iter_content(chunk_size=(DOWNLOAD_CHUNK_SIZE))):

                    if not chunk:
                        continue

                    downloaded.extend(chunk)

                    if (len(downloaded) > MAX_DOCUMENT_BYTES):
                        raise InvalidDocumentError("CV vượt quá giới hạn dung lượng cho phép.")

        except InvalidDocumentError:
            raise

        except requests.RequestException as exc:
            raise DocumentDownloadError("Không thể tải CV từ Cloudinary.") from exc

        content = bytes(downloaded)

        self._validate_content(content)

        return (content, content_type,)

    @staticmethod
    def _validate_content(
        content: bytes,
    ) -> None:
        if not content:
            raise EmptyDocumentError(
                "CV tải về không có nội dung."
            )

        if len(content) > MAX_DOCUMENT_BYTES:
            raise InvalidDocumentError("CV vượt quá giới hạn dung lượng cho phép.")

    def parse_bytes(self, content: bytes, extension: str) -> ParsedDocument:
        self._validate_content(content)

        extension = extension.lower()

        parser = self._parsers.get(extension)

        if parser is None:
            raise UnsupportedDocumentError(f"Chưa hỗ trợ định dạng CV: {extension or 'không xác định'}.")

        text = parser.parse(content)

        return ParsedDocument(
            content=content,
            text=text,
            extension=extension,
        )